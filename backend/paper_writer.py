import io
import re
from datetime import date
import asyncio
from docx import Document
import agent_wrapper

class PaperWriterAgent:
    """Agent orchestrating paper draft generation based on paper_generation and peer_review skills."""
    def __init__(self):
        self.gen_agent = agent_wrapper.AntigravityAgent("paper_generation.md")
        self.review_agent = agent_wrapper.AntigravityAgent("peer_review.md")

    async def generate_draft(self, topic: str, doc_type: str = "Research Paper", outline: str = "") -> dict:
        prompt = f"Document Type: {doc_type}\nTopic: {topic}\nOutline:\n{outline}"
        gen_res = await self.gen_agent.run(prompt)
        draft_text = gen_res.get("text", "")
        
        # Peer review
        review_res = await self.review_agent.run(draft_text)
        review_text = review_res.get("text", "")

        return {
            "status": "success",
            "doc_type": doc_type,
            "topic": topic,
            "draft": draft_text,
            "review": review_text,
            "word_count": len(re.findall(r"\b[\w'-]+\b", draft_text))
        }

def make_docx(text: str, title: str = "Para Paper V2 Document") -> bytes:
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(f"Generated/Edited on: {date.today().isoformat()}")
    document.add_heading("Content", level=1)
    for paragraph in text.splitlines():
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
